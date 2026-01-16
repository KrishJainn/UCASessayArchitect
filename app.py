import streamlit as st
import os
import io
import backend
from time import sleep
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Page Config ---
st.set_page_config(page_title="InfoYoung India - College Architect", page_icon="🎓", layout="wide")

# --- Session State Management ---
# --- Session State Management ---
# Permanent State Variables (Data)
if "step" not in st.session_state: st.session_state.step = 1
if "target_course" not in st.session_state: st.session_state.target_course = ""
if "super_curriculars" not in st.session_state: st.session_state.super_curriculars = ""
if "work_experience" not in st.session_state: st.session_state.work_experience = ""
if "student_story" not in st.session_state: st.session_state.student_story = ""
if "generated_essay" not in st.session_state: st.session_state.generated_essay = ""
if "cv_text" not in st.session_state: st.session_state.cv_text = ""

# VERSION TRACER (To help user see if Cloud Updated)
st.sidebar.info("🚀 **v4.1: Deep Brain & Anti-Lag** \n(If you see this, you have the latest code)")

# --- Helper Functions for Navigation & Persistence ---
def next_step():
    st.session_state.step += 1
def prev_step():
    st.session_state.step -= 1
def set_step(step):
    st.session_state.step = step

def save_step_1():
    # Save input widgets to persistent state
    st.session_state.target_course = st.session_state.input_target_course
    st.session_state.super_curriculars = st.session_state.input_super_curriculars
    st.session_state.work_experience = st.session_state.input_work_experience
    next_step()

def save_step_2():
    st.session_state.student_story = st.session_state.input_student_story
    next_step()

# --- CSS Injection (Premium Theme) ---
def inject_premium_css():
    st.markdown("""
        <style>
        /* Global Background */
        .stApp {
            background-color: #F3F4F6;
        }
        
        /* Hide Default Header/Footer */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        
        /* The Card Container (Glassmorphism) */
        .block-container {
            background-color: #FFFFFF;
            border-radius: 12px;
            padding: 3rem !important;
            margin-top: 2rem;
            margin-bottom: 2rem;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
            max_width: 900px;
        }
        
        /* Typography */
        h1, h2, h3 {
            color: #1A3A6D !important;
            font-family: 'Helvetica Neue', sans-serif;
        }
        
        /* Text Area & Input Styling */
        .stTextInput input, .stTextArea textarea {
            background-color: #FFFFFF;
            border: 1px solid #E5E7EB;
            border-radius: 8px;
            padding: 1rem;
        }
        .stTextInput input:focus, .stTextArea textarea:focus {
            border-color: #FAA21C;
            box-shadow: 0 0 0 2px rgba(250, 162, 28, 0.2);
        }
        
        /* Buttons */
        div.stButton > button {
            background-color: #FFFFFF;
            color: #1A3A6D;
            border: 2px solid #1A3A6D;
            border-radius: 8px;
            font-weight: 600;
            padding: 0.5rem 2rem;
            transition: all 0.2s;
        }
        div.stButton > button:hover {
            background-color: #F0F9FF;
            border-color: #1A3A6D;
            color: #1A3A6D;
        }
        
        /* Primary (Generate) Button */
        div.stButton > button[kind="primary"] {
            background-color: #FAA21C;
            border-color: #FAA21C;
            color: #1A3A6D;
        }
        div.stButton > button[kind="primary"]:hover {
            background-color: #e5941b;
            border-color: #e5941b;
            color: #14111C;
        }
        
        </style>
    """, unsafe_allow_html=True)

inject_premium_css()

# --- Sidebar Navigation ---
with st.sidebar:
    st.image("logo.png", width="stretch") 
    st.markdown("---")
    
    app_mode = st.radio("Navigation", ["🎓 UCAS Personal Statement", "⚙️ Admin: Train Brain"])
    
    st.markdown("---")
    
    # Show brain stats
    essay_count = backend.get_essay_count()
    st.caption(f"🧠 Brain Strength: Trained on {essay_count} chunks")
    
    if app_mode == "🎓 UCAS Personal Statement":
        st.markdown("### Step Progress")
        progress = (st.session_state.step / 3)
        st.progress(progress)
        st.caption(f"Step {st.session_state.step} of 3")
        
        if st.button("🔄 Reset Form"):
            for key in ["target_course", "super_curriculars", "work_experience", "student_story", "generated_essay"]:
                if key in st.session_state:
                     del st.session_state[key]
            st.session_state.step = 1
            st.rerun()

# --- Page: Admin (Train Brain) ---
if app_mode == "⚙️ Admin: Train Brain":
    st.title("⚙️ Admin: Knowledge Base")
    st.markdown("Upload successful UCAS Personal Statements to train the AI on the British academic style.")
    
    # Upload Section
    st.subheader("📤 Upload Essays")
    uploaded_files = st.file_uploader("Upload PDF or Word Essays", type=["pdf", "docx"], accept_multiple_files=True)
    
    if st.button("Process & Save to Brain"):
        if uploaded_files:
            if not os.path.exists("pdfs"): os.makedirs("pdfs")
            
            # CRITICAL: Clear cache before ingesting to ensure fresh DB connection
            st.cache_resource.clear()
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for i, uploaded_file in enumerate(uploaded_files):
                file_path = os.path.join("pdfs", uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                status_text.text(f"Ingesting {uploaded_file.name}...")
                result = backend.ingest_essay(file_path)
                st.write(f"  → {result}")  # Show result for debugging
                progress_bar.progress((i + 1) / len(uploaded_files))
            
            status_text.text("Training Complete!")
            st.success("New exemplars added to the persistent brain.")
            
            # AUTO-LEARN: Trigger analysis after new uploads
            status_text.text("🧠 Auto-analyzing all essays to learn patterns...")
            with st.spinner("Learning from your essays..."):
                analysis_result = backend.analyze_all_essays()
                if "error" not in analysis_result:
                    st.success("✅ Brain updated with new patterns!")
                else:
                    st.warning(f"Analysis skipped: {analysis_result.get('error')}")
            
            sleep(1)
            st.rerun()
        else:
            st.warning("Please upload at least one PDF.")
    
    st.markdown("---")
    
    # Global Learning Section
    st.subheader("🧠 Global Learning")
    st.markdown("Analyze ALL uploaded essays to learn their structure and style patterns.")
    
    if st.button("🔬 Analyze All & Build Blueprint", type="primary"):
        with st.spinner("Analyzing entire corpus... This may take a minute."):
            result = backend.analyze_all_essays()
            
            if "error" in result:
                st.error(f"Analysis failed: {result['error']}")
            else:
                st.success("✅ Brain Blueprint Created!")
                
                # Display learned structure
                st.subheader("📊 Learned Structure Blueprint")
                blueprint = result.get("Structure_Blueprint", {})
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Q1 (Why)", f"{blueprint.get('Q1_percentage', 20)}%")
                with col2:
                    st.metric("Q2 (Academic)", f"{blueprint.get('Q2_percentage', 25)}%")
                with col3:
                    st.metric("Q3 (Outside)", f"{blueprint.get('Q3_percentage', 55)}%")
                
                st.info(f"📝 Notes: {blueprint.get('notes', 'N/A')}")
                
                # Display Style Bible
                st.subheader("📖 Learned Style Bible")
                for rule in result.get("Style_Bible", []):
                    st.write(f"• {rule}")
                
                # Display Section Tones
                st.subheader("🎯 Section Tones")
                tones = result.get("Section_Tone", {})
                st.write(f"**Q1:** {tones.get('Q1_tone', 'N/A')}")
                st.write(f"**Q2:** {tones.get('Q2_tone', 'N/A')}")
                st.write(f"**Q3:** {tones.get('Q3_tone', 'N/A')}")
                
                # Anti-patterns
                if result.get("Anti_Patterns"):
                    st.subheader("🚫 Anti-Patterns (Things to Avoid)")
                    for ap in result.get("Anti_Patterns", []):
                        st.write(f"• {ap}")
    
    # Show current brain config if exists
    brain_config = backend.load_brain_config()
    if brain_config:
        st.markdown("---")
        st.caption(f"📅 Last analyzed: {brain_config.get('_metadata', {}).get('analysis_date', 'Unknown')}")
        st.caption(f"📊 Chunks analyzed: {brain_config.get('_metadata', {}).get('analyzed_chunks', 'Unknown')}")
    
    st.markdown("---")
    
    # Synthetic Cloner Section
    st.subheader("🧬 Synthetic Training Data Generator")
    st.markdown("Upload your BEST essay and clone it into multiple subjects to multiply your training data.")
    
    exemplar_file = st.file_uploader("Upload Best Essay (TXT or PDF)", type=["txt", "pdf"], key="exemplar_upload")
    
    col1, col2 = st.columns(2)
    with col1:
        subjects_input = st.text_area(
            "Subjects to Clone (one per line)",
            value="Computer Science\nLaw\nMedicine\nMechanical Engineering\nPhysics",
            height=150
        )
    with col2:
        st.info("The cloner will rewrite your essay for each subject while preserving:\n- Sentence structure\n- Scene-card logic\n- Emotional arc")
    
    if st.button("🧬 Generate Synthetic Training Data", type="secondary"):
        if exemplar_file:
            # Read the exemplar
            if exemplar_file.name.endswith('.txt'):
                exemplar_text = exemplar_file.read().decode('utf-8')
            else:
                # For PDF, save and use PyPDFLoader
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    tmp.write(exemplar_file.getbuffer())
                    tmp_path = tmp.name
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(tmp_path)
                pages = loader.load()
                exemplar_text = "\n".join([p.page_content for p in pages])
                os.unlink(tmp_path)
            
            subjects = [s.strip() for s in subjects_input.strip().split('\n') if s.strip()]
            
            with st.spinner(f"Generating {len(subjects)} synthetic essays... This may take 2-3 minutes."):
                results = backend.generate_synthetic_essays(exemplar_text, subjects)
                saved_files = backend.save_synthetic_essays(results)
            
            st.success(f"✅ Generated {len(saved_files)} synthetic essays!")
            
            # Show results
            for subject, essay in results:
                if not essay.startswith("Error"):
                    with st.expander(f"📄 {subject}"):
                        st.text_area("", essay[:2000] + "..." if len(essay) > 2000 else essay, height=200, key=f"syn_{subject}")
            
            st.info(f"Saved to: synthetic_essays/ folder. Use 'Process & Save to Brain' to ingest them.")
        else:
            st.warning("Please upload an exemplar essay first.")

# --- Page: Write Essay (Student Wizard) ---
elif app_mode == "🎓 UCAS Personal Statement":
    
    # Wizard Step 1: Academic Profile
    if st.session_state.step == 1:
        st.title("🧬 Step 1: Academic Profile")
        st.markdown("Define your academic interests. UCAS focuses on **super-curriculars** (what you've done outside class to explore your subject).")
        
        st.text_input("Target Course", 
                     key="input_target_course", 
                     value=st.session_state.target_course,
                     placeholder="e.g. Computer Science, Law, Economics")
        
        st.text_area("Super-Curriculars (Books, Papers, Lectures, Projects)", 
                     key="input_super_curriculars", 
                     value=st.session_state.super_curriculars, 
                     height=150,
                     placeholder="e.g. Read 'Thinking, Fast and Slow', Completed CS50, Researched Quantum Algorithms...")
        
        st.text_area("Work Experience / Relevant Skills", 
                     key="input_work_experience",
                     value=st.session_state.work_experience,
                     height=100,
                     placeholder="e.g. Internship at Fintech startup, Python (Advanced), Debate Team Captain")

        # Optional CV Upload
        st.markdown("---")
        with st.expander("📄 Optional: Upload Your CV (helps AI know you better)", expanded=False):
            st.caption("Upload your CV/Resume and we'll extract your experiences automatically.")
            uploaded_cv = st.file_uploader("Upload CV (PDF or Word)", type=["pdf", "docx"], key="cv_uploader")
            
            if uploaded_cv:
                try:
                    if uploaded_cv.name.endswith('.pdf'):
                        import tempfile
                        from langchain_community.document_loaders import PyPDFLoader
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                            tmp.write(uploaded_cv.getbuffer())
                            tmp_path = tmp.name
                        loader = PyPDFLoader(tmp_path)
                        pages = loader.load()
                        cv_text = "\n".join([p.page_content for p in pages])
                        os.unlink(tmp_path)
                    else:
                        import docx2txt
                        cv_text = docx2txt.process(uploaded_cv)
                    
                    st.session_state.cv_text = cv_text
                    st.success(f"✅ CV loaded! ({len(cv_text)} characters extracted)")
                    
                    with st.expander("Preview extracted text"):
                        st.text(cv_text[:1000] + ("..." if len(cv_text) > 1000 else ""))
                except Exception as e:
                    st.error(f"Error reading CV: {e}")
            
            if st.session_state.cv_text:
                st.info("✓ CV data will be used to enrich your personal statement.")

        st.markdown("---")
        st.button("Next: The Motivation 👉", on_click=save_step_1)

    # Wizard Step 2: The Motivation
    elif st.session_state.step == 2:
        st.title("✍️ Step 2: The Motivation")
        st.markdown("Draft your core narrative. Why this subject? What was the 'spark'? Keep it direct.")
        
        st.text_area("Why do you want to study this? (The 'Spark')", 
                    key="input_student_story", 
                    value=st.session_state.student_story,
                    height=400, 
                    placeholder="My interest in Economics began when I observed...")
        
        st.markdown("---")
        col1, col2 = st.columns([1, 1])
        with col1:
            st.button("👈 Back", on_click=prev_step)
        with col2:
            st.button("Next: The Architect 👉", on_click=save_step_2)

    # Wizard Step 3: The Architect
    elif st.session_state.step == 3:
        st.title("🏗️ Step 3: The Architect")
        st.markdown("Ready to draft. UCAS Limit: **4,000 Characters**.")
        
        col1, col2 = st.columns(2)
        with col1:
            st.info(f"**Course:** {st.session_state.target_course}")
        with col2:
            st.info("**Style:** British Academic (Formal)")
        
        if st.button("🚀 Generate UCAS 2026 Statement", type="primary"):
            if not st.session_state.student_story:
                st.error("Please provide your motivation in Step 2.")
            else:
                with st.spinner("🤖 Phoenix Generator thinking (this may take 30s)..."):
                    # 1. Prepare User Profile (include CV if available)
                    cv_section = ""
                    if st.session_state.cv_text:
                        cv_section = f"\n                    CV/Resume Details: {st.session_state.cv_text[:2000]}"
                    
                    user_profile = f"""
                    Target Course: {st.session_state.target_course}
                    Motivation: {st.session_state.student_story}
                    Super-Curriculars: {st.session_state.super_curriculars}
                    Work Experience: {st.session_state.work_experience}{cv_section}
                    """
                    
                    # 2. STAGE 1: Retrieve ALL essays from the brain (corpus analysis)
                    vectorstore = backend.get_vectorstore()
                    essay_count = backend.get_essay_count()
                    
                    # CRITICAL: Check if brain is empty BEFORE searching
                    if essay_count == 0:
                        st.error("🧠 **Brain is empty!** Please go to 'Admin: Train Brain' and upload essay PDFs first.")
                        st.stop()
                    
                    # Get a broad sample first (up to 50 chunks to show "Whole Brain" analysis)
                    all_essays = vectorstore.similarity_search("personal statement motivation academic", k=min(essay_count, 50))
                    corpus_text = "\n\n".join([doc.page_content for doc in all_essays])
                    
                    # 3. STAGE 2: Get the BEST exemplars matched to THIS student's profile
                    # Combine course + motivation for targeted retrieval
                    search_query = f"{st.session_state.target_course} {st.session_state.student_story[:500]}"
                    # PERFORMANCE UPDATE: Gemini Flash has 1M context. Increasing to 30 for max style transfer.
                    best_exemplars = vectorstore.similarity_search(search_query, k=min(essay_count, 30))
                    retrieved_exemplars = "\n\n---EXEMPLAR---\n\n".join([doc.page_content for doc in best_exemplars])
                    
                    # 4. Load Brain Config (Rules)
                    brain_config = backend.load_brain_config() or {}
                    
                    # 5. DEBUG: Show what the AI is reading
                    with st.expander("🧠 See what the AI is reading (Style Bible)"):
                        st.info(f"**Stage 1:** Analyzed {len(all_essays)} chunks ({len(corpus_text)} chars) from entire corpus")
                        st.info(f"**Stage 2:** Selected {len(best_exemplars)} best-matched exemplars ({len(retrieved_exemplars)} chars)")
                        if len(retrieved_exemplars) < 50:
                            st.error("⚠️ WARNING: No text retrieved! Upload essays to Admin first.")
                        else:
                            st.code(retrieved_exemplars[:1500] + ("..." if len(retrieved_exemplars) > 1500 else ""))
                    
                    # 5. Call Phoenix Generator
                    try:
                        result_json = backend.generate_separated_essay(user_profile, retrieved_exemplars, brain_config)
                        
                        if "error" in result_json:
                            st.error(f"Generation Error: {result_json['error']}")
                        else:
                            st.session_state.generated_essay_json = result_json
                            st.session_state.generated_essay = "draft_generated"
                            
                            # SHOW THE AI THINKING (Visible Chain of Thought)
                            if "analysis_log" in result_json:
                                with st.expander("🧠 AI Analysis & Strategy (Step-by-Step)", expanded=True):
                                    st.write(result_json["analysis_log"])
                    except Exception as e:
                         st.error(f"Frontend Error: {e}")

        if st.session_state.get("generated_essay") == "draft_generated" and st.session_state.get("generated_essay_json"):
            result_json = st.session_state.generated_essay_json
            
            # Get sections (backend returns q1_answer, q2_answer, q3_answer)
            q1 = result_json.get('q1_answer', '')
            q2 = result_json.get('q2_answer', '')
            q3 = result_json.get('q3_answer', '')
            
            # DEFENSIVE: Ensure values are strings (API might return lists/tuples)
            if isinstance(q1, (list, tuple)):
                q1 = ' '.join(str(x) for x in q1)
            if isinstance(q2, (list, tuple)):
                q2 = ' '.join(str(x) for x in q2)
            if isinstance(q3, (list, tuple)):
                q3 = ' '.join(str(x) for x in q3)
            
            # Convert to string if needed
            q1 = str(q1) if q1 else ''
            q2 = str(q2) if q2 else ''
            q3 = str(q3) if q3 else ''
            
            # Calculate character counts
            q1_chars = len(q1)
            q2_chars = len(q2)
            q3_chars = len(q3)
            total_chars = q1_chars + q2_chars + q3_chars
            
            st.success("✅ Drafting Complete! Phoenix Engine Active.")
            
            # --- QUALITY GATE CHECK ---
            passed, issues, score = backend.quality_gate(f"{q1}\n{q2}\n{q3}")
            if not passed:
                st.error(f"⚠️ **AI DETECTOR WARNING (Score: {score}/100)**: logic detected robotic patterns.")
                for issue in issues:
                    st.warning(f"• {issue}")
            else:
                st.info(f"🛡️ **Antigravity Quality Check Passed** (Score: {score}/100)")
            # ---------------------------
            
            # Character count metrics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Q1", f"{q1_chars} chars")
            with col2:
                st.metric("Q2", f"{q2_chars} chars")
            with col3:
                st.metric("Q3", f"{q3_chars} chars")
            with col4:
                delta = total_chars - 4000
                st.metric("Total", f"{total_chars}/4000", delta=f"{delta:+d}" if delta != 0 else "✓ Exact", delta_color="inverse")
            
            st.markdown("---")
            
            # Match the Next.js frontend styling
            st.markdown("""
            <style>
            .ucas-question {
                font-size: 0.75rem;
                font-weight: 700;
                color: #9CA3AF;
                text-transform: uppercase;
                letter-spacing: 0.1em;
                margin-bottom: 0.75rem;
            }
            .ucas-answer {
                font-family: Georgia, serif;
                color: #1F2937;
                line-height: 1.75;
                padding: 1rem;
                background-color: #F9FAFB;
                border-radius: 0.5rem;
                border: 1px solid #E5E7EB;
            }
            </style>
            """, unsafe_allow_html=True)
            
            # Q1
            st.markdown('<p class="ucas-question">1️⃣ Why do you want to study this course or subject?</p>', unsafe_allow_html=True)
            st.markdown(f'<div class="ucas-answer">{q1}</div>', unsafe_allow_html=True)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Q2
            st.markdown('<p class="ucas-question">2️⃣ How have your qualifications and studies prepared you for this course?</p>', unsafe_allow_html=True)
            st.markdown(f'<div class="ucas-answer">{q2}</div>', unsafe_allow_html=True)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Q3
            st.markdown('<p class="ucas-question">3️⃣ What else have you done to prepare for this course?</p>', unsafe_allow_html=True)
            st.markdown(f'<div class="ucas-answer">{q3}</div>', unsafe_allow_html=True)

            # Combine for download
            full_essay = f"{q1}\n\n{q2}\n\n{q3}"
            
            # Generate Word Document
            doc = Document()
            
            # Title
            title = doc.add_heading('UCAS Personal Statement', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add character count info
            info = doc.add_paragraph(f"Total Characters: {total_chars}/4000")
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.runs[0]
            info_run.font.size = Pt(10)
            info_run.font.italic = True
            
            doc.add_paragraph()
            
            # Q1 Section
            doc.add_heading('Why do you want to study this course?', level=2)
            q1_para = doc.add_paragraph(q1)
            q1_para.paragraph_format.space_after = Pt(12)
            
            # Q2 Section
            doc.add_heading('How have your studies prepared you?', level=2)
            q2_para = doc.add_paragraph(q2)
            q2_para.paragraph_format.space_after = Pt(12)
            
            # Q3 Section
            doc.add_heading('What else have you done to prepare?', level=2)
            q3_para = doc.add_paragraph(q3)
            q3_para.paragraph_format.space_after = Pt(12)
            
            # Save to bytes for download
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            # Download buttons
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "� Download as Word",
                    doc_bytes.getvalue(),
                    file_name="UCAS_Personal_Statement.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col2:
                st.download_button("📝 Download as Text", full_essay, file_name="UCAS_Personal_Statement.txt")
                
        st.markdown("---")
        st.button("👈 Back to Motivation", on_click=prev_step)
