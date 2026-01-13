
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import backend
import os
import shutil
import tempfile
from typing import Optional

app = FastAPI(title="College Architect API")

# Enable CORS for Next.js frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Local development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Pydantic Models for Request/Response ---

class UserProfile(BaseModel):
    target_course: str
    motivation: str
    super_curriculars: str
    work_experience: str
    cv_text: Optional[str] = ""

class GenerateRequest(BaseModel):
    profile: UserProfile

# --- Endpoints ---

@app.get("/health")
def health_check():
    return {"status": "ok", "service": "College Architect API"}

@app.get("/")
def root():
    return {"message": "College Architect Brain is Active 🧠", "docs": "/docs"}

@app.get("/stats")
def get_stats():
    """Returns brain statistics (essay count)."""
    try:
        count = backend.get_essay_count()
        return {"essay_count": count}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze")
def trigger_analysis():
    """Triggers global corpus analysis."""
    try:
        result = backend.analyze_all_essays()
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["error"])
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate")
def generate_essay(req: GenerateRequest):
    """Generates an essay using the Phoenix engine."""
    try:
        # 1. Construct User Profile String (matching app.py logic)
        cv_section = f"\nCV/Resume Details: {req.profile.cv_text[:2000]}" if req.profile.cv_text else ""
        
        full_profile_str = f"""
        Target Course: {req.profile.target_course}
        Motivation: {req.profile.motivation}
        Super-Curriculars: {req.profile.super_curriculars}
        Work Experience: {req.profile.work_experience}{cv_section}
        """

        # 2. Retrieve Exemplars (Backend Logic)
        vectorstore = backend.get_vectorstore()
        
        # Broad Search first
        # Note: We can't easily access the vectorstore count directly here without backend help, 
        # but backend.get_vectorstore() returns the object.
        # Let's rely on backend logic or replicate it slightly.
        # Ideally, move the retrieval logic into a clean function in backend.py, 
        # but for now we will implement the logic here to match app.py's flow.
        
        essay_count = backend.get_essay_count()
        if essay_count == 0:
            raise HTTPException(status_code=400, detail="Brain is empty. Please upload essays first.")

        # Stage 2 Retrieval (Targeted)
        search_query = f"{req.profile.target_course} {req.profile.motivation[:500]}"
        best_exemplars = vectorstore.similarity_search(search_query, k=5)
        retrieved_exemplars = "\n\n---EXEMPLAR---\n\n".join([doc.page_content for doc in best_exemplars])

        # 3. Load Config
        brain_config = backend.load_brain_config() or {}

        # 4. Generate
        result_json = backend.generate_separated_essay(full_profile_str, retrieved_exemplars, brain_config)
        
        if "error" in result_json:
            raise HTTPException(status_code=500, detail=result_json["error"])
            
        return result_json

    except Exception as e:
        print(f"API Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ingest")
async def ingest_file(file: UploadFile = File(...)):
    """Ingests a PDF or DOCX file into the brain."""
    try:
        # Create 'pdfs' dir if not exists (backend expects it usually, or we just need temp)
        if not os.path.exists("pdfs"):
            os.makedirs("pdfs")

        file_path = os.path.join("pdfs", file.filename)
        
        # Save uploaded file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        # Ingest
        result = backend.ingest_essay(file_path)
        
        return {"filename": file.filename, "status": "ingested", "message": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/parse-cv")
async def parse_cv(file: UploadFile = File(...)):
    """Extracts text from an uploaded CV (PDF or DOCX)."""
    try:
        if not os.path.exists("temp_cvs"):
            os.makedirs("temp_cvs")

        file_path = os.path.join("temp_cvs", file.filename)
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        text = ""
        if file.filename.lower().endswith(".pdf"):
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            text = "\n".join([p.page_content for p in pages])
        elif file.filename.lower().endswith(".docx"):
            import docx2txt
            text = docx2txt.process(file_path)
        
        # Cleanup
        os.remove(file_path)
        
        return {"text": text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

from docx import Document
from docx.shared import Pt
from fastapi.responses import FileResponse

class DownloadRequest(BaseModel):
    q1: str
    q2: str
    q3: str

@app.post("/download-docx")
def download_docx(req: DownloadRequest):
    try:
        # Create a new Document
        doc = Document()
        
        # Title
        title = doc.add_heading('UCAS Personal Statement', 0)
        title.alignment = 1  # Center
        
        # Q1
        doc.add_heading('1. Why do you want to study this course or subject?', level=2)
        p1 = doc.add_paragraph(req.q1)
        p1.paragraph_format.space_after = Pt(12)
        
        # Q2
        doc.add_heading('2. How have your qualifications and studies prepared you?', level=2)
        p2 = doc.add_paragraph(req.q2)
        p2.paragraph_format.space_after = Pt(12)
        
        # Q3
        doc.add_heading('3. What else have you done to prepare for this course?', level=2)
        p3 = doc.add_paragraph(req.q3)
        p3.paragraph_format.space_after = Pt(12)
        
        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        temp_file.close()
        
        return FileResponse(
            path=temp_file.name, 
            filename="Personal_Statement_Draft.docx",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
         raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
